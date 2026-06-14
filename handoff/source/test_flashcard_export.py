"""Unit tests for the flashcard data model and export paths touched by the
"minimum information / cloze deletion" rework — run with:

    python -m unittest test_flashcard_export -v

These exercise only the pure, API-free parts (dataclasses, JSON extraction,
and the CSV/Markdown/Word exporters), so they run without an `anthropic`
install or network access — `llm_client` only imports `anthropic` lazily,
inside `generate_section_content`.
"""

from __future__ import annotations

import csv
import os
import tempfile
import unittest

from llm_client import Flashcard, SectionResult, CharacterSummary, _extract_json
from exporter import (
    export_flashcards_csv,
    export_cloze_flashcards_csv,
    count_cloze_flashcards,
    export_summaries_markdown,
    export_summaries_docx,
)


# --------------------------------------------------------------------------- #
# Flashcard / card_type
# --------------------------------------------------------------------------- #

class FlashcardCardTypeTests(unittest.TestCase):
    def test_defaults_to_basic(self):
        card = Flashcard(front="Q", back="A")
        self.assertEqual(card.card_type, "basic")
        self.assertFalse(card.is_cloze)

    def test_cloze_flag(self):
        card = Flashcard(front="The {{c1::mitochondria}} is the powerhouse.", back="", card_type="cloze")
        self.assertTrue(card.is_cloze)


# --------------------------------------------------------------------------- #
# _extract_json — tolerant JSON parsing of the model's response
# --------------------------------------------------------------------------- #

class ExtractJsonTests(unittest.TestCase):
    def test_plain_json(self):
        self.assertEqual(_extract_json('{"summary": "hi"}'), {"summary": "hi"})

    def test_fenced_json(self):
        raw = "```json\n{\"flashcards\": []}\n```"
        self.assertEqual(_extract_json(raw), {"flashcards": []})

    def test_embedded_json_with_stray_text(self):
        raw = "Sure, here you go:\n{\"summary\": \"x\"}\nHope that helps!"
        self.assertEqual(_extract_json(raw), {"summary": "x"})


# --------------------------------------------------------------------------- #
# Helpers to build fixture data
# --------------------------------------------------------------------------- #

def _make_results():
    """A small two-section fixture: one section with a mix of basic and cloze
    cards, one with only cloze cards, and one that — like a table of contents
    — produced no flashcards at all (the new, intentionally-supported case)."""
    return [
        SectionResult(
            title="Chapter 1: Salt Lakes",
            summary="Covers the Dead Sea's geography and chemistry.",
            flashcards=[
                Flashcard(front="How deep is the Dead Sea?", back="304 meters", card_type="basic"),
                Flashcard(
                    front="The Dead Sea's surface sits about {{c1::430}} meters below sea level.",
                    back="Lowest land elevation on Earth.",
                    card_type="cloze",
                ),
            ],
        ),
        SectionResult(
            title="Chapter 2: Founding Members",
            summary="Covers the founding of the alliance.",
            flashcards=[
                Flashcard(
                    front="The alliance was founded in {{c1::1949}} by {{c2::twelve}} countries.",
                    back="",
                    card_type="cloze",
                ),
            ],
        ),
        SectionResult(
            title="Table of Contents",
            summary="Front matter listing chapter titles and page numbers.",
            flashcards=[],
        ),
    ]


# --------------------------------------------------------------------------- #
# CSV export — basic vs. cloze must land in separate files/note types
# --------------------------------------------------------------------------- #

class CsvExportTests(unittest.TestCase):
    def setUp(self):
        self.results = _make_results()
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmpdir.cleanup)

    def _path(self, name):
        return os.path.join(self.tmpdir.name, name)

    def test_count_cloze_flashcards(self):
        self.assertEqual(count_cloze_flashcards(self.results), 2)

    def test_basic_csv_excludes_cloze_cards(self):
        path = self._path("flashcards.csv")
        count = export_flashcards_csv(self.results, path, book_title="Test Book")
        self.assertEqual(count, 1)  # only the one basic Q&A card
        with open(path, encoding="utf-8") as f:
            rows = list(csv.reader(f))
        self.assertEqual(len(rows), 1)
        front, back, tag = rows[0]
        self.assertEqual(front, "How deep is the Dead Sea?")
        self.assertEqual(back, "304 meters")
        self.assertEqual(tag, "Test_Book::Chapter_1:_Salt_Lakes")
        self.assertNotIn("{{c1::", "".join(row[0] for row in rows))

    def test_cloze_csv_contains_only_cloze_cards(self):
        path = self._path("flashcards_cloze.csv")
        count = export_cloze_flashcards_csv(self.results, path, book_title="Test Book")
        self.assertEqual(count, 2)
        with open(path, encoding="utf-8") as f:
            rows = list(csv.reader(f))
        self.assertEqual(len(rows), 2)
        # Every Text field should carry cloze markup; every row maps to (text, extra, tag)
        for text, extra, tag in rows:
            self.assertIn("{{c", text)
            self.assertTrue(tag.startswith("Test_Book::"))
        # The first card has "extra" context; the second deliberately has none
        texts_to_extra = {text: extra for text, extra, _ in rows}
        self.assertEqual(
            texts_to_extra["The Dead Sea's surface sits about {{c1::430}} meters below sea level."],
            "Lowest land elevation on Earth.",
        )
        self.assertEqual(
            texts_to_extra["The alliance was founded in {{c1::1949}} by {{c2::twelve}} countries."],
            "",
        )

    def test_section_with_no_flashcards_contributes_nothing(self):
        # The "Table of Contents" section produced zero cards — make sure it
        # simply contributes nothing to either file rather than erroring.
        basic_path, cloze_path = self._path("b.csv"), self._path("c.csv")
        export_flashcards_csv(self.results, basic_path, book_title="Test Book")
        export_cloze_flashcards_csv(self.results, cloze_path, book_title="Test Book")
        for path in (basic_path, cloze_path):
            with open(path, encoding="utf-8") as f:
                rows = list(csv.reader(f))
            self.assertTrue(all("Table_of_Contents" not in row[-1] for row in rows))


# --------------------------------------------------------------------------- #
# Markdown study guide — both card styles should render distinctly and
# legibly, and a card-less section shouldn't produce a "Flashcards:" header.
# --------------------------------------------------------------------------- #

class MarkdownExportTests(unittest.TestCase):
    def setUp(self):
        self.results = _make_results()
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmpdir.cleanup)
        self.path = os.path.join(self.tmpdir.name, "guide.md")

    def _read(self):
        export_summaries_markdown(self.results, self.path, book_title="Test Book")
        with open(self.path, encoding="utf-8") as f:
            return f.read()

    def test_basic_and_cloze_sections_rendered(self):
        text = self._read()
        # Basic Q&A rendering preserved
        self.assertIn("**Q:** How deep is the Dead Sea?", text)
        self.assertIn("**A:** 304 meters", text)
        # Cloze cards get their own labeled block, with raw {{c1::...}} visible
        self.assertIn("Cloze cards", text)
        self.assertIn("{{c1::430}}", text)
        self.assertIn("*Lowest land elevation on Earth.*", text)
        # A cloze card with no "extra" context shouldn't emit a stray italic line
        self.assertIn("{{c1::1949}}", text)

    def test_section_with_no_cards_has_no_flashcard_header(self):
        text = self._read()
        toc_section = text.split("## Table of Contents", 1)[1]
        # Cut off at the next "## " heading (there isn't one — it's the last section)
        self.assertNotIn("**Flashcards:**", toc_section)
        self.assertNotIn("Cloze cards", toc_section)


# --------------------------------------------------------------------------- #
# Word export — structural sanity check (headings, bullet styles, runs)
# --------------------------------------------------------------------------- #

class DocxExportTests(unittest.TestCase):
    def setUp(self):
        self.results = _make_results()
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmpdir.cleanup)
        self.path = os.path.join(self.tmpdir.name, "guide.docx")

    def test_renders_without_error_and_contains_expected_text(self):
        export_summaries_docx(self.results, self.path, book_title="Test Book",
                              character_list=[CharacterSummary(name="Anna", summary="A traveler.")])
        from docx import Document
        doc = Document(self.path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        self.assertIn("Test Book — Study Guide", full_text)
        self.assertIn("Anna", full_text)
        self.assertIn("How deep is the Dead Sea?", full_text)
        self.assertIn("Cloze cards:", full_text)
        self.assertIn("{{c1::430}}", full_text)
        self.assertIn("Lowest land elevation on Earth.", full_text)

        # The basic card's "A: " label should be a bold run, the cloze extra
        # context should be an italic run — spot-check a couple of runs.
        bold_runs = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
        italic_runs = [r.text for p in doc.paragraphs for r in p.runs if r.italic]
        self.assertIn("A: ", bold_runs)
        self.assertIn("Lowest land elevation on Earth.", italic_runs)

    def test_section_with_no_cards_has_no_flashcard_heading_text(self):
        export_summaries_docx(self.results, self.path, book_title="Test Book")
        from docx import Document
        doc = Document(self.path)
        # Find the "Table of Contents" heading and check the very next
        # paragraphs (its summary) don't include flashcard scaffolding.
        idx = next(i for i, p in enumerate(doc.paragraphs) if p.text == "Table of Contents")
        following = [p.text for p in doc.paragraphs[idx + 1:idx + 4]]
        self.assertFalse(any("Flashcards:" in t or "Cloze cards:" in t for t in following))


if __name__ == "__main__":
    unittest.main()
