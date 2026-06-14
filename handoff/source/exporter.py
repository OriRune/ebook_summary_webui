"""Export generated results to Anki-importable CSV (in two flavors — ordinary
Basic-style cards and cloze-deletion cards, since Anki needs each in its own
note type), a Markdown study guide, and a formatted Word (.docx) study guide."""

from __future__ import annotations

import csv
from llm_client import SectionResult, CharacterSummary


def _section_tag(result: SectionResult, book_title: str) -> str:
    return (book_title or "ebook").replace(" ", "_") + "::" + result.title.replace(" ", "_")


def count_cloze_flashcards(results: list[SectionResult]) -> int:
    """How many of the generated cards are cloze-deletion style. Used by the
    GUI to decide whether a second export pass (see export_cloze_flashcards_csv)
    is needed at all."""
    return sum(1 for r in results for c in r.flashcards if c.is_cloze)


def export_flashcards_csv(results: list[SectionResult], path: str, book_title: str = "") -> int:
    """Write a CSV with columns: front, back, tags — ready for Anki's
    File > Import (map column 1 -> Front, column 2 -> Back, column 3 -> Tags)
    using the **Basic** note type. Returns the number of cards written.

    Cloze-deletion cards are deliberately skipped here — Anki only renders
    `{{c1::...}}` markup as a fill-in-the-blank when the note uses its
    **Cloze** note type, which expects different fields (Text / Back Extra).
    Mixing the two into one Basic import would just show the curly-brace
    syntax as literal text. See export_cloze_flashcards_csv for those."""
    count = 0
    with open(path, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        for result in results:
            tag = _section_tag(result, book_title)
            for card in result.flashcards:
                if card.is_cloze:
                    continue
                writer.writerow([card.front, card.back, tag])
                count += 1
    return count


def export_cloze_flashcards_csv(results: list[SectionResult], path: str, book_title: str = "") -> int:
    """Write a CSV of cloze-deletion cards with columns: text, back extra,
    tags — ready for Anki's File > Import using the **Cloze** note type (map
    column 1 -> Text, column 2 -> Back Extra, column 3 -> Tags). Anki parses
    the `{{c1::...}}` markup in the Text field itself and builds the
    fill-in-the-blank card automatically. Returns the number of cards written
    (0 if the book produced no cloze cards — callers can skip this file
    entirely in that case; see count_cloze_flashcards)."""
    count = 0
    with open(path, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        for result in results:
            tag = _section_tag(result, book_title)
            for card in result.flashcards:
                if not card.is_cloze:
                    continue
                writer.writerow([card.front, card.back, tag])
                count += 1
    return count


def export_summaries_markdown(
    results: list[SectionResult],
    path: str,
    book_title: str = "",
    character_list: list[CharacterSummary] | None = None,
) -> None:
    # Determine whether all successful sections used the same model so we can
    # show it once in the header (clean) vs per-section (mixed-model run).
    models_used = [r.model_used for r in results if not r.error and r.model_used]
    unique_models = sorted(set(models_used))

    lines = []
    if book_title:
        lines.append(f"# {book_title} — Study Guide\n")
    if len(unique_models) == 1:
        lines.append(f"_Generated with: {unique_models[0]}_\n")
    elif len(unique_models) > 1:
        lines.append(f"_Generated with multiple models: {', '.join(unique_models)}_\n")
    if character_list:
        lines.append("## Main Characters\n")
        for character in character_list:
            lines.append(f"**{character.name}**")
            lines.append(f"{character.summary}\n")
    for result in results:
        lines.append(f"## {result.title}\n")
        if result.error:
            lines.append(f"*Generation failed: {result.error}*\n")
            continue
        if len(unique_models) > 1 and result.model_used:
            lines.append(f"_Model: {result.model_used}_\n")
        if result.summary:
            lines.append(f"{result.summary}\n")
        if result.flashcards:
            basic_cards = [c for c in result.flashcards if not c.is_cloze]
            cloze_cards = [c for c in result.flashcards if c.is_cloze]
            if basic_cards:
                lines.append("**Flashcards:**\n")
                for card in basic_cards:
                    lines.append(f"- **Q:** {card.front}")
                    lines.append(f"  **A:** {card.back}")
                lines.append("")
            if cloze_cards:
                lines.append(
                    "**Cloze cards** (Anki fill-in-the-blank style — the "
                    "`{{c1::...}}` portion is the part you'd be asked to recall):\n"
                )
                for card in cloze_cards:
                    lines.append(f"- {card.front}")
                    if card.back:
                        lines.append(f"  *{card.back}*")
                lines.append("")
        if result.discussion_questions:
            lines.append("**Discussion questions:**\n")
            for question in result.discussion_questions:
                lines.append(f"- {question}")
            lines.append("")
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def export_summaries_docx(
    results: list[SectionResult],
    path: str,
    book_title: str = "",
    character_list: list[CharacterSummary] | None = None,
) -> None:
    """Write the same study-guide content as export_summaries_markdown, but as
    a formatted Word document. Markdown is great as a portable, diffable
    source — but as the thing you actually sit down and read, print, or
    annotate from, raw `##`/`**` markup gets in the way. A .docx renders
    headings, bold labels, and bullet lists as real formatting, and opens
    directly in Word, Google Docs, or LibreOffice."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "Exporting as a Word document requires the 'python-docx' package. "
            "Install it with: pip install python-docx"
        ) from e

    # Same model-display logic as the markdown exporter.
    models_used = [r.model_used for r in results if not r.error and r.model_used]
    unique_models = sorted(set(models_used))

    doc = Document()

    if book_title:
        doc.add_heading(f"{book_title} — Study Guide", level=0)

    if len(unique_models) == 1:
        model_p = doc.add_paragraph()
        model_p.add_run(f"Generated with: {unique_models[0]}").italic = True
    elif len(unique_models) > 1:
        model_p = doc.add_paragraph()
        model_p.add_run(f"Generated with multiple models: {', '.join(unique_models)}").italic = True

    if character_list:
        doc.add_heading("Main Characters", level=1)
        for character in character_list:
            name_p = doc.add_paragraph()
            name_p.add_run(character.name).bold = True
            doc.add_paragraph(character.summary)

    for result in results:
        doc.add_heading(result.title, level=1)
        if result.error:
            err_p = doc.add_paragraph()
            err_p.add_run(f"Generation failed: {result.error}").italic = True
            continue
        if len(unique_models) > 1 and result.model_used:
            model_p = doc.add_paragraph()
            model_p.add_run(f"Model: {result.model_used}").italic = True
        if result.summary:
            doc.add_paragraph(result.summary)
        if result.flashcards:
            basic_cards = [c for c in result.flashcards if not c.is_cloze]
            cloze_cards = [c for c in result.flashcards if c.is_cloze]
            if basic_cards:
                label_p = doc.add_paragraph()
                label_p.add_run("Flashcards:").bold = True
                for card in basic_cards:
                    q_p = doc.add_paragraph(style="List Bullet")
                    q_p.add_run("Q: ").bold = True
                    q_p.add_run(card.front)
                    a_p = doc.add_paragraph(style="List Bullet")
                    a_p.add_run("A: ").bold = True
                    a_p.add_run(card.back)
            if cloze_cards:
                label_p = doc.add_paragraph()
                label_p.add_run("Cloze cards:").bold = True
                note_p = doc.add_paragraph()
                note_p.add_run(
                    "(Anki fill-in-the-blank style — the {{c1::...}} portion "
                    "is the part you'd be asked to recall)"
                ).italic = True
                for card in cloze_cards:
                    doc.add_paragraph(card.front, style="List Bullet")
                    if card.back:
                        extra_p = doc.add_paragraph()
                        extra_p.add_run(card.back).italic = True
        if result.discussion_questions:
            label_p = doc.add_paragraph()
            label_p.add_run("Discussion questions:").bold = True
            for question in result.discussion_questions:
                doc.add_paragraph(question, style="List Bullet")

    doc.save(path)
